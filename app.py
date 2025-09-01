from flask import Flask, render_template, request, send_file
from PIL import Image
from pdf2image import convert_from_path
import pdfplumber
import docx
from fpdf import FPDF
import os

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
CONVERT_FOLDER = 'converted'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERT_FOLDER, exist_ok=True)

POPPLER_PATH = r"C:\poppler\Library\bin"  # Vérifie ce chemin

@app.route('/')
def index():
    return render_template('index.html')


# 1️⃣ PNG -> PDF
@app.route('/png-to-pdf', methods=['GET', 'POST'])
def png_to_pdf():
    if request.method == 'POST':
        try:
            files = request.files.getlist('images')
            if not files or files[0].filename == '':
                return "Aucun fichier sélectionné."
            images = []
            for file in files:
                file_path = os.path.join(UPLOAD_FOLDER, file.filename)
                file.save(file_path)
                img = Image.open(file_path)
                if img.mode == "RGBA":
                    img = img.convert("RGB")
                images.append(img)
            pdf_path = os.path.join(CONVERT_FOLDER, 'converted.pdf')
            images[0].save(pdf_path, save_all=True, append_images=images[1:])
            return send_file(pdf_path, as_attachment=True)
        except Exception as e:
            return f"Erreur PNG -> PDF : {e}"
    return render_template('png-to-pdf.html')


# 2️⃣ PDF -> PNG
@app.route('/pdf-to-png', methods=['GET', 'POST'])
def pdf_to_png():
    if request.method == 'POST':
        try:
            file = request.files['file']
            if not file or file.filename == '':
                return "Aucun fichier PDF sélectionné."
            pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(pdf_path)
            # Conversion PDF -> PNG
            images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)
            png_files = []
            for i, img in enumerate(images):
                png_filename = f"{os.path.splitext(file.filename)[0]}_page_{i+1}.png"
                png_path = os.path.join(CONVERT_FOLDER, png_filename)
                img.save(png_path, 'PNG')
                png_files.append(png_path)
            return f"Conversion terminée ! {len(png_files)} pages converties en PNG."
        except Exception as e:
            return f"Erreur PDF -> PNG : {e}\n\nVérifie que Poppler est installé et que le PDF n’est pas protégé ou corrompu."
    return render_template('pdf-to-png.html')


# 3️⃣ Word -> PDF (UTF-8)
@app.route('/word-to-pdf', methods=['GET', 'POST'])
def word_to_pdf():
    if request.method == 'POST':
        try:
            file = request.files['file']
            if not file or file.filename == '':
                return "Aucun fichier Word sélectionné."
            file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(file_path)
            doc = docx.Document(file_path)
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", '', 12)
            for para in doc.paragraphs:
                # Encodage UTF-8 sûr : remplace les caractères non supportés
                text = para.text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, text)
            pdf_path = os.path.join(CONVERT_FOLDER, os.path.splitext(file.filename)[0]+'.pdf')
            pdf.output(pdf_path)
            return send_file(pdf_path, as_attachment=True)
        except Exception as e:
            return f"Erreur Word -> PDF : {e}"
    return render_template('word-to-pdf.html')


# 4️⃣ PDF -> Word
@app.route('/pdf-to-word', methods=['GET', 'POST'])
def pdf_to_word():
    if request.method == 'POST':
        try:
            file = request.files['file']
            if not file or file.filename == '':
                return "Aucun fichier PDF sélectionné."
            file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(file_path)
            word_path = os.path.join(CONVERT_FOLDER, os.path.splitext(file.filename)[0]+'.docx')
            doc = docx.Document()
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
            doc.save(word_path)
            return send_file(word_path, as_attachment=True)
        except Exception as e:
            return f"Erreur PDF -> Word : {e}"
    return render_template('pdf-to-word.html')


if __name__ == '__main__':
    app.run(debug=True)
